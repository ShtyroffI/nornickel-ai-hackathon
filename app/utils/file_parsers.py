# utils/file_parsers.py

import logging
import re
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple, Union

logger = logging.getLogger(__name__)


def parse_docx(file_path: str) -> str:
    """
    Извлекает текст из DOCX-файла (абзацы + таблицы).
    Текст из таблиц добавляется с пометкой [TABLE].
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed. Run: pip install python-docx")
        return ""

    doc = Document(file_path)
    text_parts = []

    # Абзацы
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text.strip())

    # Таблицы
    for table in doc.tables:
        table_text = []
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_text.append(" | ".join(row_text))
        if table_text:
            text_parts.append("[TABLE] " + " ; ".join(table_text))

    return "\n".join(text_parts)


def extract_tables_docx(file_path: str) -> List[List[List[str]]]:
    """
    Извлекает таблицы из DOCX как список таблиц,
    каждая таблица — список строк, каждая строка — список ячеек (строки).
    """
    try:
        from docx import Document
    except ImportError:
        logger.error("python-docx not installed")
        return []

    doc = Document(file_path)
    all_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        all_tables.append(rows)
    return all_tables


# === PDF парсер (текстовый + OCR) ===

def parse_pdf(file_path: str, use_ocr_fallback: bool = True) -> str:
    """
    Извлекает текст из PDF.
    Сначала пытается через pdfplumber (текстовый PDF).
    Если текста мало или ошибка, и use_ocr_fallback=True — запускает OCR.
    """
    text = _parse_pdf_textual(file_path)
    if text and len(text.strip()) > 50:  # Если текста достаточно
        return text
    if use_ocr_fallback:
        logger.info(f"PDF {file_path} seems scanned or has little text, trying OCR...")
        return _parse_pdf_ocr(file_path)
    return text


def _parse_pdf_textual(file_path: str) -> str:
    """Извлекает текст из текстового PDF через pdfplumber."""
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed. Run: pip install pdfplumber")
        return ""

    try:
        with pdfplumber.open(file_path) as pdf:
            text_parts = []
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text.strip())
                # Извлечение таблиц (текст из таблиц тоже добавляется)
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        row_text = " | ".join([str(cell or "").strip() for cell in row])
                        if row_text:
                            text_parts.append("[TABLE] " + row_text)
            return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"Error parsing PDF with pdfplumber: {e}")
        return ""


def _parse_pdf_ocr(file_path: str) -> str:
    """
    Распознаёт текст через OCR (pytesseract + pdf2image).
    Требует установленных Tesseract и poppler-utils.
    """
    try:
        from pdf2image import convert_from_path
        import pytesseract
    except ImportError as e:
        logger.error(f"OCR libraries not installed: {e}. Install: pip install pdf2image pytesseract")
        return ""

    try:
        images = convert_from_path(file_path, dpi=300)
        text_parts = []
        for _, img in enumerate(images):
            # Язык: русский + английский
            page_text = pytesseract.image_to_string(img, lang="rus+eng")
            if page_text.strip():
                text_parts.append(page_text.strip())
        return "\n".join(text_parts)
    except Exception as e:
        logger.error(f"OCR failed: {e}")
        return ""


def extract_tables_pdf(file_path: str) -> List[List[List[str]]]:
    """
    Извлекает таблицы из PDF (только из текстовых, с помощью pdfplumber).
    """
    try:
        import pdfplumber
    except ImportError:
        logger.error("pdfplumber not installed")
        return []

    tables = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_tables = page.extract_tables()
                for table in page_tables:
                    # Преобразуем None в пустую строку
                    clean_table = [
                        [str(cell or "").strip() for cell in row]
                        for row in table
                    ]
                    if clean_table:
                        tables.append(clean_table)
    except Exception as e:
        logger.error(f"Error extracting tables from PDF: {e}")
    return tables


# === TXT парсер ===

def parse_txt(file_path: str) -> str:
    """Просто читает текстовый файл."""
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        # Попробуем другую кодировку
        with open(file_path, "r", encoding="cp1251") as f:
            return f.read()


# === Главный диспетчер ===

def parse_file(file_path: Union[str, Path]) -> str:
    """
    Основная функция: определяет тип файла по расширению и извлекает текст.
    Поддерживаемые расширения: .docx, .pdf, .txt (регистронезависимо).
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return parse_docx(str(file_path))
    elif suffix == ".pdf":
        return parse_pdf(str(file_path), use_ocr_fallback=True)
    elif suffix == ".txt":
        return parse_txt(str(file_path))
    else:
        logger.warning(f"Unsupported file type: {suffix}")
        return ""


def extract_tables(file_path: Union[str, Path]) -> List[List[List[str]]]:
    """
    Извлекает таблицы из файла (DOCX или PDF) в виде списка таблиц.
    Для TXT возвращает пустой список.
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".docx":
        return extract_tables_docx(str(file_path))
    elif suffix == ".pdf":
        return extract_tables_pdf(str(file_path))
    else:
        return []


def parse_file_with_tables(file_path: Union[str, Path]) -> Tuple[str, List[List[List[str]]]]:
    """
    Возвращает (текст, список_таблиц) для файла.
    Удобно, если нужно раздельно обрабатывать текст и структуру таблиц.
    """
    text = parse_file(file_path)
    tables = extract_tables(file_path)
    return text, tables


def clean_text(text: str) -> str:
    """Простая очистка: удаляет лишние пробелы, пустые строки."""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    return "\n".join(lines)
